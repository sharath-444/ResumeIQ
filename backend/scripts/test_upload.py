import requests
import os
from docx import Document

BASE_URL = "http://127.0.0.1:5000"

def create_dummy_docx():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('johndoe@example.com | 555-0199')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, Flask, JavaScript, SQL, Git')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at Test Corp')
    doc.save("test_resume.docx")
    return "test_resume.docx"

def test_upload():
    session = requests.Session()
    
    # 1. Login
    print("Logging in...")
    login_res = session.post(f"{BASE_URL}/auth/login", json={"username": "admin", "password": "password123"})
    if login_res.status_code != 200:
        print(f"Login failed: {login_res.text}")
        return

    # 2. Upload
    filename = create_dummy_docx()
    print(f"Uploading {filename}...")
    
    with open(filename, 'rb') as f:
        files = {'resume': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'role': 'Backend Developer'}
        
        try:
            res = session.post(f"{BASE_URL}/upload", files=files, data=data)
            print(f"Status Code: {res.status_code}")
            if res.status_code == 200:
                print("Upload Success!")
                print(res.json())
            else:
                print("Upload Failed")
                print(res.text)
        except Exception as e:
            print(f"Exception: {e}")

if __name__ == "__main__":
    test_upload()
